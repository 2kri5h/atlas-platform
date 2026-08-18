import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

NAVY_HEX    = "1F3864"
LBLUE_HEX   = "D9E1F2"
BORDER_HEX  = "8EA9DB"

NAVY   = RGBColor(31,  56,  100)
DARK   = RGBColor(33,  37,   41)
GRAY   = RGBColor(89,  89,   89)
WHITE  = RGBColor(255, 255, 255)

# ── helpers ─────────────────────────────────────────────────────────────────

def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def cell_pad(cell, top=80, bot=80, lft=110, rgt=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for name, val in [('top',top),('bottom',bot),('left',lft),('right',rgt)]:
        n = OxmlElement(f'w:{name}')
        n.set(qn('w:w'), str(val)); n.set(qn('w:type'), 'dxa')
        m.append(n)
    tcPr.append(m)

def tbl_borders(table, hex_color=BORDER_HEX, sz="4"):
    table._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top    w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'<w:left   w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'<w:right  w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:color="{hex_color}"/>'
        f'</w:tblBorders>'))

def rule(p, hex_color=NAVY_HEX, sz="8"):
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{hex_color}"/>'
        f'</w:pBdr>'))

def fld(run, instruction):
    r = run._r
    for t, x in [("begin",None),("separate",None),("end",None)]:
        if t == "separate":
            r.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> {instruction} </w:instrText>'))
        r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="{t}"/>'))

# ── document builders ────────────────────────────────────────────────────────

def run(p, text, bold=False, italic=False, size=10.5, color=None):
    r = p.add_run(text)
    r.font.name  = 'Times New Roman'
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    r.font.color.rgb = color or DARK
    return r

def body(doc, text, bold_prefix="", after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run(p, bold_prefix, bold=True)
    run(p, text)
    return p

def bullet(doc, text, bold_prefix="", after=3, indent=False):
    style = 'List Bullet 2' if indent else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        run(p, bold_prefix, bold=True)
    run(p, text)
    return p

def heading(doc, text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.keep_with_next = True
    rule(p, NAVY_HEX, "8")
    run(p, text, bold=True, size=13, color=NAVY)
    return p

def subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.keep_with_next = True
    run(p, text, bold=True, size=11, color=NAVY)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    run(p, text, italic=True, size=9.5, color=GRAY)
    return p

def spacer(doc, before=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(0)

# ── main ─────────────────────────────────────────────────────────────────────

def create_document():
    doc = docx.Document()

    # page setup
    sec = doc.sections[0]
    sec.page_width   = Inches(8.27)
    sec.page_height  = Inches(11.69)
    sec.top_margin   = Inches(0.65)
    sec.bottom_margin= Inches(0.65)
    sec.left_margin  = Inches(0.75)
    sec.right_margin = Inches(0.75)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = DARK

    # ensure List Bullet 2 exists (for nested bullets)
    try:
        doc.styles['List Bullet 2']
    except KeyError:
        lb2 = doc.styles.add_style('List Bullet 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        lb2.base_style = doc.styles['List Bullet']

    # ── header ────────────────────────────────────────────────────────────
    hdr = sec.header
    ht  = hdr.add_table(1, 2, Inches(6.77))
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER

    lc = ht.cell(0, 0); lc.width = Inches(3.0)
    hp = lc.paragraphs[0]
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)
    logo = 'c:/Users/krish/Development/ITSP/Team-atlas-ITSP/logo_1.png'
    if os.path.exists(logo):
        hp.add_run().add_picture(logo, height=Inches(0.55))

    rc = ht.cell(0, 1); rc.width = Inches(3.77)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    run(rp, "Institute Technical Council\n", bold=True,  size=11, color=NAVY)
    run(rp, "Institute Technical Summer Projects",        italic=True, size=9.5, color=NAVY)

    hl = hdr.add_paragraph()
    hl.paragraph_format.space_before = Pt(4)
    hl.paragraph_format.space_after  = Pt(0)
    rule(hl, NAVY_HEX, "6")

    # ── footer ─────────────────────────────────────────────────────────────
    ftr = sec.footer
    fp  = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page "); r1.font.name='Times New Roman'; r1.font.size=Pt(9); r1.font.color.rgb=GRAY
    fld(r1, "PAGE")
    r2 = fp.add_run(" of ");  r2.font.name='Times New Roman'; r2.font.size=Pt(9); r2.font.color.rgb=GRAY
    fld(r2, "NUMPAGES")

    # ── title block ────────────────────────────────────────────────────────
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(10)
    pt.paragraph_format.space_after  = Pt(2)
    run(pt, "Institute Technical Summer Projects (ITSP)", bold=True, size=17, color=NAVY)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after  = Pt(10)
    run(ps, "Project Reference Documentation", size=12.5, color=GRAY)

    # ── metadata table ─────────────────────────────────────────────────────
    mt = doc.add_table(4, 2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_borders(mt)
    for lbl, val in [
        ("Team Name", ""),
        ("Project Title", ""),
        ("Domain / Category", "Web Development, AI/ML, Productivity & Mental Health Tools"),
        ("Mentor(s)", ""),
    ]:
        i = [("Team Name",""),("Project Title",""),("Domain / Category","Web Development, AI/ML, Productivity & Mental Health Tools"),("Mentor(s)","")].index((lbl, val))
        lc = mt.rows[i].cells[0]; lc.width = Inches(2.0)
        shade(lc, LBLUE_HEX); cell_pad(lc)
        lp = lc.paragraphs[0]; lp.paragraph_format.space_before=Pt(0); lp.paragraph_format.space_after=Pt(0)
        run(lp, lbl, bold=True, size=10, color=NAVY)

        vc = mt.rows[i].cells[1]; vc.width = Inches(4.77); cell_pad(vc)
        vp = vc.paragraphs[0]; vp.paragraph_format.space_before=Pt(0); vp.paragraph_format.space_after=Pt(0)
        if val:
            run(vp, val, italic=True, size=10, color=GRAY)

    spacer(doc, 4)

    # ==================================================================
    # SECTION 1 - TEAM MEMBERS
    # ==================================================================
    heading(doc, "1. Team Members")
    note(doc, "List every team member with their official roll number (format: 25BXXXX).")

    tt = doc.add_table(5, 4)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_borders(tt)
    hdrs  = ["S.No.", "Name", "Roll Number", "Role / Responsibility"]
    cwids = [Inches(0.65), Inches(2.2), Inches(1.6), Inches(2.32)]

    for i, h in enumerate(hdrs):
        c = tt.rows[0].cells[i]; c.width = cwids[i]
        shade(c, NAVY_HEX); cell_pad(c, 90, 90, 100, 100)
        p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        run(p, h, bold=True, size=10, color=WHITE)

    for ri in range(1, 5):
        for ci in range(4):
            c = tt.rows[ri].cells[ci]; c.width = cwids[ci]; cell_pad(c, 70, 70, 100, 100)
            p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            if ci == 0: run(p, str(ri), size=10)
            elif ci == 2: run(p, "25B_____", size=10, color=GRAY)

    # ==================================================================
    # SECTION 2 - PROJECT DETAILS
    # ==================================================================
    heading(doc, "2. Project Details")

    # ── 2.1 Problem Statement ──────────────────────────────────────────
    subheading(doc, "2.1  Problem Statement")

    body(doc,
        "Students at IIT Bombay juggle rigorous academics, placements, extracurriculars, and"
        " personal well-being simultaneously. Several structural gaps make this harder than it needs to be:",
        after=3)

    for line in [
        "Academic timetables, lab slots, assignment deadlines, and announcements are scattered"
        " across institute webmail, course portals, and external apps - leading to clashes and missed deadlines.",
        "There is no automated way to convert email announcements into calendar entries without manual effort.",
        "Junior students lack a structured, searchable archive of senior academic and career experiences.",
        "Privacy concerns prevent students from openly asking sensitive course or placement questions on public forums.",
        "No existing tool scientifically tracks student workload against burnout risk indicators.",
    ]:
        bullet(doc, line, after=2)

    spacer(doc, 4)

    # ── 2.2 Solution Proposed ─────────────────────────────────────────
    subheading(doc, "2.2  Solution Proposed")

    body(doc,
        "The ITSP Student Productivity and Well-being Platform is a full-stack web application"
        " purpose-built for IIT Bombay students. It brings together six integrated modules into one unified interface:",
        after=4)

    modules = [
        ("Smart Planner:",
         "Interactive 0-24 hour calendar across daily, weekly, and monthly views. Supports event"
         " location, course code tagging, deadlines, meeting links, importance flags, activity type"
         " classification (academic / working-hour / personal), and checklist notes."),
        ("Webmail Ingestion Pipeline:",
         "Connects to institute email via IMAP with AES-256 encrypted token storage. Strips HTML noise,"
         " generates 2-3 sentence email digests, and extracts events (exams, deadlines, club activities)"
         " as structured JSON using the Gemini LLM. A SHA-256 deduplication engine prevents redundant"
         " API calls for mass broadcast emails, cutting compute costs by 80-90%."),
        ("AI Academic Mentor:",
         "A Gemini-powered chat assistant with persistent multi-session context. Connected to the student's"
         " planner and resource database to generate personalised study schedules and goal summaries."),
        ("Burnout Prediction System:",
         "A Gradient Boosting ML model (trained on 150,000 student profiles) evaluates six daily behavioral"
         " inputs - Study Hours, Sleep Duration, Social Support, Physical Activity, Screen Time, and CGPA -"
         " to classify burnout risk as Low, Medium, or High with tailored health recommendations."),
        ("Resource Library:",
         "A community-driven repository with domain filters (Software Engineering, AI/ML, Data Science),"
         " keyword search, upvoting, bookmarking, private uploads, and a first-visit interactive onboarding"
         " tour built with React Joyride."),
        ("Senior Journeys and Anonymous Q&A Portal:",
         "A structured archive of senior student experiences (academics, placements, research) paired with"
         " a privacy-first, thread-based Q&A portal for anonymous junior-senior interaction."),
    ]

    for bold_lbl, desc in modules:
        bullet(doc, desc, bold_prefix=bold_lbl + "  ", after=4)

    body(doc,
        "What sets this apart from generic tools: automated campus email parsing with SHA-256 deduplication,"
        " an NIH-research-backed ML burnout model, a context-aware AI mentor connected to live planner data,"
        " and safe anonymous mentorship channels built specifically for an IIT Bombay context.",
        after=4)

    # ── 2.3 Work Done So Far ──────────────────────────────────────────
    subheading(doc, "2.3  Work / Simulations Done Till Now")

    body(doc, "The following has been completed across research, architecture, simulation, and prototype testing:",
         after=3)

    bullet(doc,
        "Conducted a user study among IITB undergraduates to map scheduling friction points, webmail overload,"
        " and information barriers in junior-senior guidance. Reviewed NIH research linking study hours to"
        " emotional exhaustion and sleep/social support to academic resilience.",
        bold_prefix="Literature review / design study completed:  ", after=4)

    bullet(doc,
        "Built a complete full-stack architecture using FastAPI, SQLAlchemy, and Supabase/SQLite for the"
        " backend REST API, and React + Vite + TypeScript for the frontend.",
        bold_prefix="CAD models / simulations built (tool used, key parameters):  ", after=2)
    bullet(doc,
        "Trained a Gradient Boosting model (website_gradient_model.pkl) on 150,000 synthetic student"
        " profiles. Key parameters: Study Hours (0-16 hrs), Sleep Hours (0-12 hrs), Social Support"
        " Score (1-10), Physical Activity (0-5 hrs), Screen Time (0-16 hrs), CGPA (0-10).",
        indent=True, after=2)
    bullet(doc,
        "Designed relational database schemas and Supabase pgvector embedding models for semantic"
        " relevance scoring of broadcast emails.",
        indent=True, after=4)

    bullet(doc,
        "Smart Planner: verified 0-24 hour grid across daily, weekly, and monthly views. Confirmed"
        " event location persistence, course code tagging (CS101), deadline tracking, meeting links,"
        " and checklist notes.",
        bold_prefix="Prototypes fabricated and tests conducted:  ", after=2)
    bullet(doc,
        "Webmail Pipeline: verified IMAP connection, AES-256 token encryption, HTML sanitization,"
        " and Gemini LLM JSON event extraction. SHA-256 deduplication engine tested across simulated"
        " mass emails - confirmed 80-90% reduction in duplicate API calls.",
        indent=True, after=2)
    bullet(doc,
        "Resource Library: implemented domain filtering, upvoting, bookmarking, and React Joyride"
        " onboarding tour with localStorage state persistence.",
        indent=True, after=2)
    bullet(doc,
        "AI Mentor and Anonymous Portal: tested persistent multi-session chat with Google GenAI SDK"
        " and verified thread-based anonymous post creation with privacy masking.",
        indent=True, after=4)

    bullet(doc,
        "SHA-256 email deduplication eliminates redundant LLM calls for campus-wide broadcasts,"
        " reducing server overhead significantly.",
        bold_prefix="Key results or design decisions arrived at:  ", after=2)
    bullet(doc,
        "Gradient Boosting model reliably classifies Low, Medium, and High burnout risk tiers"
        " across varied student behavioral profiles.",
        indent=True, after=2)
    bullet(doc,
        "Settled on 1-click action cards ([Add to Calendar], [Set Reminder], [Ask AI Mentor])"
        " as the primary interaction pattern to minimise manual friction.",
        indent=True, after=4)

    # ── 2.4 Timeline for Remaining Work ───────────────────────────────
    subheading(doc, "2.4  Timeline for Remaining Work")

    body(doc, "Four phases remain before final submission:", after=3)

    phases = [
        ("Phase 1 - Focus-First UI/UX Overhaul:",
         "Redesign the interface into a student dashboard with three clear zones: Today's Priorities,"
         " Upcoming Deadlines, and Actionable Emails. Replace list views with interactive action cards"
         " ([Add to Calendar], [Set Reminder], [Ask AI Mentor])."),
        ("Phase 2 - End-to-End Security and Vector Recommender:",
         "Enforce Supabase Row-Level Security (auth.uid() = user_id) on all tables and apply pgcrypto"
         " column-level encryption to sensitive fields. Build pgvector semantic embeddings to compute"
         " email Relevance Scores (High / Medium / Low) based on student major, year, and club memberships."),
        ("Phase 3 - AI Study Breakdowns and Extended Features:",
         "Enable automated study plan generation for upcoming exams. Integrate Google Drive resource"
         " link attachments into planner tasks if time permits."),
        ("Phase 4 - Supabase Cloud Deployment and Full Data Migration:",
         "Migrate the backend from local SQLite to Supabase (PostgreSQL). Tables to migrate include:"
         " user accounts and profiles, AI chatbot conversation history, parsed webmail summaries,"
         " planner calendar entries, burnout assessment logs, Senior Journeys posts, and Anonymous"
         " Q&A threads. Frontend (React + Vite) to be deployed on Vercel. FastAPI backend to be"
         " containerised and deployed on Railway or Render. End-to-end acceptance testing on the"
         " live environment before final submission."),
    ]

    for bold_lbl, desc in phases:
        bullet(doc, desc, bold_prefix=bold_lbl + "  ", after=4)

    # ── 2.5 Expected Outcome ──────────────────────────────────────────
    subheading(doc, "2.5  Expected Outcome / Deliverable")

    body(doc,
        "The final deliverable is a fully cloud-deployed, production-ready web application for IIT Bombay"
        " students. All application data (user profiles, AI chatbot history, planner events, webmail"
        " summaries, burnout logs, anonymous posts, and resource bookmarks) will be stored in Supabase"
        " with Row-Level Security, pgcrypto encryption, and pgvector embeddings.",
        after=3)

    body(doc, "Core features demonstrated at submission:", after=3)

    for item in [
        "AI-powered Webmail Ingestion with SHA-256 deduplication and 1-click calendar population",
        "Interactive Smart Planner (daily, weekly, monthly views) with full event metadata support",
        "GenAI Academic Mentor with persistent multi-session chat history and study plan generation",
        "ML Burnout Prediction (Gradient Boosting, 6 behavioral inputs, Low/Medium/High classification)",
        "Domain-filtered Resource Library with React Joyride interactive onboarding tour",
        "Senior Journeys knowledge archive and Anonymous Q&A Portal",
    ]:
        bullet(doc, item, after=2)

    out = 'c:/Users/krish/Development/ITSP/Team-atlas-ITSP/ITSP_Project_Reference_Documentation.docx'
    doc.save(out)
    print(f'Saved: {out}')

if __name__ == '__main__':
    create_document()
